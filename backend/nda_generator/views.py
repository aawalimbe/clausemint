from rest_framework.decorators import api_view
from rest_framework.response import Response
from rest_framework import status
from django.views.decorators.csrf import csrf_exempt
from django.http import HttpResponse, StreamingHttpResponse
from django.conf import settings
from django.utils import timezone
import json
import os
import logging
from datetime import datetime
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from core.ai_service import ai_service
from core.jurisdictions import get_default_jurisdiction
from .utils import load_nda_prompt

# Configure logging
logger = logging.getLogger(__name__)


def generate_nda_streaming(messages, max_tokens=2000, temperature=0.3):
    """Generate NDA content using streaming"""
    try:
        # Use the AI service with streaming
        response = ai_service._generate_mistral_response(messages, max_tokens=max_tokens, temperature=temperature)
        return response
    except Exception as e:
        logger.error(f"Streaming generation error: {e}")
        raise e


@api_view(['POST'])
@csrf_exempt
def generate_nda_streaming_view(request):
    """Generate NDA with streaming response"""
    try:
        data = json.loads(request.body)
        
        # Extract NDA parameters
        party_a = data.get('party_a', '')
        party_b = data.get('party_b', '')
        party_c = data.get('party_c', '')
        nda_type = data.get('nda_type', 'two-way')
        purpose = data.get('purpose', '')
        confidentiality_period = data.get('confidentiality_period', '2 years')
        jurisdiction = data.get('jurisdiction', get_default_jurisdiction())
        
        # Validate required fields
        if not party_a or not party_b or not purpose:
            logger.warning(f"NDA generation failed: Missing required fields - Party A: {party_a}, Party B: {party_b}, Purpose: {purpose}")
            return Response({
                'error': 'Party A, Party B, and Purpose are required fields'
            }, status=status.HTTP_400_BAD_REQUEST)
        
        # Log NDA generation attempt
        logger.info(f"Starting streaming NDA generation for {party_a} and {party_b}, Type: {nda_type}, Jurisdiction: {jurisdiction}")
        
        # Load the NDA prompt template
        prompt_template = load_nda_prompt()
        
        # Format the prompt with user inputs
        formatted_prompt = prompt_template.format(
            party_a=party_a,
            party_b=party_b,
            party_c=party_c if party_c else 'N/A',
            nda_type=nda_type,
            purpose=purpose,
            confidentiality_period=confidentiality_period,
            jurisdiction=jurisdiction
        )
        
        # Generate NDA using AI service
        messages = [
            {"role": "system", "content": "You are a legal document generator specializing in NDA agreements. Generate professional, legally sound documents."},
            {"role": "user", "content": formatted_prompt}
        ]
        
        def generate_stream():
            """Generate streaming response"""
            try:
                # Send initial status
                yield f"data: {json.dumps({'status': 'started', 'message': 'Starting NDA generation...'})}\n\n"
                
                # Generate content using streaming
                content = generate_nda_streaming(messages, max_tokens=2000, temperature=0.3)
                
                # Send the complete content
                yield f"data: {json.dumps({'status': 'completed', 'content': content, 'parameters': {
                    'party_a': party_a,
                    'party_b': party_b,
                    'party_c': party_c,
                    'nda_type': nda_type,
                    'purpose': purpose,
                    'confidentiality_period': confidentiality_period,
                    'jurisdiction': jurisdiction
                }})}\n\n"
                
                # Log successful generation
                logger.info(f"Streaming NDA generated successfully for {party_a} and {party_b}")
                
            except Exception as e:
                error_msg = f"Streaming NDA generation failed: {str(e)}"
                logger.error(error_msg, exc_info=True)
                yield f"data: {json.dumps({'status': 'error', 'error': str(e)})}\n\n"
        
        # Return streaming response
        response = StreamingHttpResponse(
            generate_stream(),
            content_type='text/event-stream'
        )
        response['Cache-Control'] = 'no-cache'
        response['X-Accel-Buffering'] = 'no'
        return response
        
    except Exception as e:
        # Log error details
        error_msg = f"Streaming NDA generation failed: {str(e)}"
        logger.error(error_msg, exc_info=True)
        
        # Write error to file
        error_log_path = os.path.join(settings.BASE_DIR, 'logs', 'nda_errors.log')
        os.makedirs(os.path.dirname(error_log_path), exist_ok=True)
        
        with open(error_log_path, 'a', encoding='utf-8') as f:
            timestamp = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
            f.write(f"[{timestamp}] {error_msg}\n")
        
        return Response({'error': 'Failed to generate NDA. Please try again.'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


@api_view(['POST'])
@csrf_exempt
def generate_nda(request):
    """Generate NDA document using AI"""
    try:
        data = json.loads(request.body)
        
        # Extract NDA parameters
        party_a = data.get('party_a', '')
        party_b = data.get('party_b', '')
        party_c = data.get('party_c', '')
        nda_type = data.get('nda_type', 'two-way')
        purpose = data.get('purpose', '')
        confidentiality_period = data.get('confidentiality_period', '2 years')
        jurisdiction = data.get('jurisdiction', get_default_jurisdiction())
        
        # Validate required fields
        if not party_a or not party_b or not purpose:
            logger.warning(f"NDA generation failed: Missing required fields - Party A: {party_a}, Party B: {party_b}, Purpose: {purpose}")
            return Response({
                'error': 'Party A, Party B, and Purpose are required fields'
            }, status=status.HTTP_400_BAD_REQUEST)
        
        # Log NDA generation attempt
        logger.info(f"Starting NDA generation for {party_a} and {party_b}, Type: {nda_type}, Jurisdiction: {jurisdiction}")
        
        # Load the NDA prompt template
        prompt_template = load_nda_prompt()
        
        # Format the prompt with user inputs
        formatted_prompt = prompt_template.format(
            party_a=party_a,
            party_b=party_b,
            party_c=party_c if party_c else 'N/A',
            nda_type=nda_type,
            purpose=purpose,
            confidentiality_period=confidentiality_period,
            jurisdiction=jurisdiction
        )
        
        # Generate NDA using AI service
        messages = [
            {"role": "system", "content": "You are a legal document generator specializing in NDA agreements. Generate professional, legally sound documents."},
            {"role": "user", "content": formatted_prompt}
        ]
        
        nda_content = ai_service.generate_response(messages, max_tokens=2000, temperature=0.3)
        
        # Log successful generation
        logger.info(f"NDA generated successfully for {party_a} and {party_b}")
        
        return Response({
            'nda_content': nda_content,
            'parameters': {
                'party_a': party_a,
                'party_b': party_b,
                'party_c': party_c,
                'nda_type': nda_type,
                'purpose': purpose,
                'confidentiality_period': confidentiality_period,
                'jurisdiction': jurisdiction
            }
        })
        
    except Exception as e:
        # Log error details
        error_msg = f"NDA generation failed: {str(e)}"
        logger.error(error_msg, exc_info=True)
        
        # Write error to file
        error_log_path = os.path.join(settings.BASE_DIR, 'logs', 'nda_errors.log')
        os.makedirs(os.path.dirname(error_log_path), exist_ok=True)
        
        with open(error_log_path, 'a', encoding='utf-8') as f:
            timestamp = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
            f.write(f"[{timestamp}] {error_msg}\n")
        
        return Response({'error': 'Failed to generate NDA. Please try again.'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


@api_view(['POST'])
@csrf_exempt
def export_nda(request):
    """Export NDA as DOCX file"""
    try:
        data = json.loads(request.body)
        nda_content = data.get('nda_content', '')
        party_a = data.get('party_a', 'Party A')
        party_b = data.get('party_b', 'Party B')
        
        if not nda_content:
            logger.warning("NDA export failed: No content provided")
            return Response({'error': 'NDA content is required'}, status=status.HTTP_400_BAD_REQUEST)
        
        # Create Word document
        doc = Document()
        
        # Set document margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # Add title
        title = doc.add_heading('NON-DISCLOSURE AGREEMENT', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add date
        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_para.add_run(f"Date: {timezone.now().strftime('%B %d, %Y')}")
        
        # Add content
        content_paragraphs = nda_content.split('\n\n')
        for para_text in content_paragraphs:
            if para_text.strip():
                # Check if it's a heading (all caps or starts with number)
                if para_text.strip().isupper() or para_text.strip()[0].isdigit():
                    doc.add_heading(para_text.strip(), level=1)
                else:
                    doc.add_paragraph(para_text.strip())
        
        # Generate filename
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f"NDA_{party_a.replace(' ', '_')}_{party_b.replace(' ', '_')}_{timestamp}.docx"
        
        # Save to media directory
        media_dir = os.path.join(settings.MEDIA_ROOT, 'nda_documents')
        os.makedirs(media_dir, exist_ok=True)
        
        file_path = os.path.join(media_dir, filename)
        doc.save(file_path)
        
        # Log successful export
        logger.info(f"NDA exported successfully: {filename}")
        
        return Response({
            'success': True,
            'filename': filename,
            'download_url': f'/media/nda_documents/{filename}',
            'message': 'NDA document generated successfully!'
        })
        
    except Exception as e:
        # Log error details
        error_msg = f"NDA export failed: {str(e)}"
        logger.error(error_msg, exc_info=True)
        
        # Write error to file
        error_log_path = os.path.join(settings.BASE_DIR, 'logs', 'nda_export_errors.log')
        os.makedirs(os.path.dirname(error_log_path), exist_ok=True)
        
        with open(error_log_path, 'a', encoding='utf-8') as f:
            timestamp = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
            f.write(f"[{timestamp}] {error_msg}\n")
        
        return Response({'error': 'Failed to export NDA document. Please try again.'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


@api_view(['GET'])
def download_nda(request, filename):
    """Download generated NDA document"""
    try:
        file_path = os.path.join(settings.MEDIA_ROOT, 'nda_documents', filename)
        
        if not os.path.exists(file_path):
            logger.warning(f"NDA download failed: File not found - {filename}")
            return Response({'error': 'File not found'}, status=status.HTTP_404_NOT_FOUND)
        
        # Log download
        logger.info(f"NDA downloaded: {filename}")
        
        # Return file for download
        with open(file_path, 'rb') as f:
            response = HttpResponse(f.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            response['Content-Disposition'] = f'attachment; filename="{filename}"'
            return response
            
    except Exception as e:
        error_msg = f"NDA download failed: {str(e)}"
        logger.error(error_msg, exc_info=True)
        return Response({'error': 'Failed to download file'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR) 